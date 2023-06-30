import docx

# doc = docx.Document('demo.docx')
# fullText = []
# for p in doc.paragraphs:
#     fullText.append(p.text)
#     print(f'Paragraph: {p.text}')
#     for r in p.runs:
#         print(f'\tRun: {r.text}')

# print('\n')
# print('\n\n'.join(fullText))

# doc.paragraphs[0].style = 'Normal'
# doc.paragraphs[1].runs[0].style = 'QuoteChar'
# doc.paragraphs[1].runs[1].bold = True
# doc.paragraphs[1].runs[2].italic = True
# doc.save('restyled.docx')

newDoc = docx.Document()
newDoc.add_heading('Header 0', 0)
newDoc.add_paragraph('hello world')
newDoc.add_heading('Header 1', 1)
secParagraph = newDoc.add_paragraph('This is second paragraph.')
newDoc.add_heading('Header 2', 2)
newDoc.add_paragraph('This is yet another paragraph.')
secParagraph.runs[-1].add_break()
secParagraph.add_run('This text is being added to second paragraph.')
newDoc.add_picture('zophie.png')
newDoc.save('HelloWorld.docx')
